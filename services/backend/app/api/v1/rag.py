from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from typing import Optional
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.database import get_db
from app.core.security.auth import get_current_user
from app.core.rag.engine import RAGEngine
from app.models.base import User, Document, UploadedFile

router = APIRouter()
rag_engine = RAGEngine()


class QueryRequest(BaseModel):
    query: str
    limit: int = 5


@router.post("/query")
async def query_knowledge(
    req: QueryRequest,
    user: User = Depends(get_current_user),
):
    results = await rag_engine.retrieve(
        query=req.query,
        user_id=user.id,
        limit=req.limit,
    )
    return {"results": results}


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    import aiofiles
    import os
    from app.config.settings import get_settings

    settings = get_settings()
    upload_dir = os.path.join(settings.STORAGE_PATH, user.id)
    os.makedirs(upload_dir, exist_ok=True)

    file_path = os.path.join(upload_dir, file.filename)
    async with aiofiles.open(file_path, "wb") as f:
        content = await file.read()
        await f.write(content)

    # Extract text based on mime type
    extracted_text = ""
    mime_type = file.content_type or "application/octet-stream"

    if "pdf" in mime_type:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            extracted_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif "text" in mime_type or "json" in mime_type or "csv" in mime_type:
        extracted_text = content.decode("utf-8", errors="ignore")
    elif "word" in mime_type or "docx" in mime_type:
        import docx
        doc = docx.Document(file_path)
        extracted_text = "\n".join(p.text for p in doc.paragraphs)

    # Save file record
    uploaded = UploadedFile(
        user_id=user.id,
        filename=file.filename,
        original_filename=file.filename,
        file_size=len(content),
        mime_type=mime_type,
        file_path=file_path,
        file_type="document",
        extracted_text=extracted_text[:50000],
    )
    db.add(uploaded)
    await db.flush()

    # Index in vector DB
    doc_record = Document(
        user_id=user.id,
        title=file.filename,
        content=extracted_text,
        file_id=uploaded.id,
    )
    db.add(doc_record)
    await db.flush()

    if extracted_text:
        chunk_count = await rag_engine.index_document(
            document_id=doc_record.id,
            content=extracted_text,
            user_id=user.id,
            metadata={"filename": file.filename},
        )
        doc_record.chunk_count = chunk_count
        doc_record.is_embedded = True
        uploaded.is_embedded = True

    return {
        "document_id": doc_record.id,
        "file_id": uploaded.id,
        "filename": file.filename,
        "chunk_count": doc_record.chunk_count,
        "message": "Document uploaded and indexed",
    }


@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from sqlalchemy import select
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.user_id == user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    await rag_engine.delete_document(document_id)
    doc.is_deleted = True

    return {"message": "Document deleted"}
